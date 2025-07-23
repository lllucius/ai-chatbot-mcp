"""
File processing utilities for document content extraction.

This module provides functions for extracting text content from various
file formats including PDF, DOCX, TXT, MD, and RTF files with
streaming support and memory optimization.

Generated on: 2025-07-14 03:18:45 UTC
Current User: lllucius
"""

import asyncio
import logging
import mimetypes
import os
from pathlib import Path
from typing import Any, AsyncIterator, Dict

import psutil

from ..core.exceptions import DocumentError, ValidationError

logger = logging.getLogger(__name__)


class FileProcessor:
    """
    Optimized file processor for extracting text content from various file formats.

    Supports PDF, DOCX, TXT, MD, RTF with streaming support and memory optimization
    for large files.
    """

    def __init__(self, max_file_size: int = 50 * 1024 * 1024, chunk_size: int = 8192):
        """
        Initialize file processor.

        Args:
            max_file_size: Maximum file size to process (50MB default)
            chunk_size: Chunk size for streaming operations (8KB default)
        """
        self.max_file_size = max_file_size
        self.chunk_size = chunk_size
        self.supported_types = {
            "pdf": self._extract_pdf,
            "docx": self._extract_docx,
            "txt": self._extract_text,
            "md": self._extract_text,
            "rtf": self._extract_rtf,
        }

    def _check_memory_usage(self) -> None:
        """Check if memory usage is too high and warn."""
        try:
            memory = psutil.virtual_memory()
            if memory.percent > 80:
                logger.warning(f"High memory usage: {memory.percent}%")
                if memory.percent > 90:
                    raise MemoryError("System memory usage too high")
        except Exception as e:
            logger.debug(f"Memory check failed: {e}")

    def _validate_file(self, file_path: str) -> None:
        """Validate file before processing."""
        path = Path(file_path)

        if not path.exists():
            raise DocumentError(f"File not found: {file_path}")

        if not path.is_file():
            raise DocumentError(f"Path is not a file: {file_path}")

        file_size = path.stat().st_size
        if file_size > self.max_file_size:
            raise DocumentError(
                f"File too large: {file_size / (1024*1024):.1f}MB > {self.max_file_size / (1024*1024):.1f}MB"
            )

        if file_size == 0:
            raise DocumentError("File is empty")

    async def extract_text(self, file_path: str, file_type: str) -> str:
        """
        Extract text content from a file with memory optimization.

        Args:
            file_path: Path to the file
            file_type: File extension/type

        Returns:
            str: Extracted text content

        Raises:
            DocumentError: If extraction fails
            ValidationError: If file type not supported
        """
        # Validate file
        self._validate_file(file_path)

        file_type = file_type.lower().lstrip(".")

        if file_type not in self.supported_types:
            raise ValidationError(f"Unsupported file type: {file_type}")

        # Check memory before processing
        self._check_memory_usage()

        try:
            extractor = self.supported_types[file_type]
            content = await extractor(file_path)

            if not content.strip():
                raise DocumentError("No text content found in file")

            return content.strip()

        except MemoryError:
            logger.error(f"Out of memory processing {file_path}")
            raise DocumentError("File too large to process - insufficient memory")
        except Exception as e:
            logger.error(f"Text extraction failed for {file_path}: {e}")
            raise DocumentError(f"Text extraction failed: {e}")

    async def extract_text_streaming(
        self, file_path: str, file_type: str
    ) -> AsyncIterator[str]:
        """
        Extract text content from a file using streaming for large files.

        Args:
            file_path: Path to the file
            file_type: File extension/type

        Yields:
            str: Text chunks as they are processed

        Raises:
            DocumentError: If extraction fails
            ValidationError: If file type not supported
        """
        # Validate file
        self._validate_file(file_path)

        file_type = file_type.lower().lstrip(".")

        if file_type not in self.supported_types:
            raise ValidationError(f"Unsupported file type: {file_type}")

        # For now, only plain text files support true streaming
        if file_type in ["txt", "md"]:
            async for chunk in self._extract_text_streaming(file_path):
                yield chunk
        else:
            # For other formats, extract all and yield in chunks
            content = await self.extract_text(file_path, file_type)
            chunk_size = 4096  # 4KB chunks
            for i in range(0, len(content), chunk_size):
                yield content[i : i + chunk_size]
                # Allow other tasks to run
                await asyncio.sleep(0)

    async def _extract_pdf(self, file_path: str) -> str:
        """Extract text from PDF file with memory optimization."""
        try:
            import PyPDF2

            text_content = []

            # Process PDF in async manner with memory checks
            await asyncio.sleep(0)  # Yield control

            with open(file_path, "rb") as file:
                pdf_reader = PyPDF2.PdfReader(file)
                total_pages = len(pdf_reader.pages)

                logger.info(f"Processing PDF with {total_pages} pages")

                for page_num in range(total_pages):
                    # Check memory usage periodically
                    if page_num % 10 == 0:
                        self._check_memory_usage()
                        await asyncio.sleep(0)  # Allow other tasks

                    try:
                        page = pdf_reader.pages[page_num]
                        page_text = page.extract_text()

                        if page_text.strip():
                            text_content.append(page_text)

                        # For very large PDFs, limit memory usage
                        if (
                            len(text_content) > 1000
                        ):  # Limit to ~1000 pages of text in memory
                            logger.warning("PDF too large, processing first 1000 pages")
                            break

                    except Exception as e:
                        logger.warning(
                            f"Failed to extract text from page {page_num}: {e}"
                        )
                        continue

            return "\n\n".join(text_content)

        except ImportError:
            # Fallback to pypdf if PyPDF2 not available
            try:
                from pypdf import PdfReader

                text_content = []
                await asyncio.sleep(0)

                with open(file_path, "rb") as file:
                    pdf_reader = PdfReader(file)
                    total_pages = len(pdf_reader.pages)

                    logger.info(f"Processing PDF with {total_pages} pages using pypdf")

                    for page_num, page in enumerate(pdf_reader.pages):
                        if page_num % 10 == 0:
                            self._check_memory_usage()
                            await asyncio.sleep(0)

                        try:
                            page_text = page.extract_text()
                            if page_text.strip():
                                text_content.append(page_text)

                            if len(text_content) > 1000:
                                logger.warning(
                                    "PDF too large, processing first 1000 pages"
                                )
                                break

                        except Exception as e:
                            logger.warning(
                                f"Failed to extract text from page {page_num}: {e}"
                            )
                            continue

                return "\n\n".join(text_content)

            except ImportError:
                raise DocumentError(
                    "PDF processing library not available. Install PyPDF2 or pypdf."
                )
        except MemoryError:
            raise DocumentError("PDF file too large to process - insufficient memory")
        except Exception as e:
            raise DocumentError(f"PDF extraction failed: {e}")

    async def _extract_docx(self, file_path: str) -> str:
        """Extract text from DOCX file with memory optimization."""
        try:
            from docx import Document

            await asyncio.sleep(0)  # Yield control

            doc = Document(file_path)
            text_content = []
            processed_items = 0

            # Extract paragraphs with memory management
            logger.info(f"Processing DOCX with {len(doc.paragraphs)} paragraphs")

            for paragraph in doc.paragraphs:
                if processed_items % 100 == 0:  # Check every 100 paragraphs
                    self._check_memory_usage()
                    await asyncio.sleep(0)

                if paragraph.text.strip():
                    text_content.append(paragraph.text)
                    processed_items += 1

                # Limit memory usage for very large documents
                if processed_items > 10000:
                    logger.warning("DOCX too large, limiting to first 10000 paragraphs")
                    break

            # Extract tables with memory management
            logger.info(f"Processing {len(doc.tables)} tables")
            table_count = 0

            for table in doc.tables:
                if table_count % 10 == 0:  # Check every 10 tables
                    self._check_memory_usage()
                    await asyncio.sleep(0)

                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content.append(" | ".join(row_text))

                table_count += 1

                # Limit table processing
                if table_count > 100:
                    logger.warning("Too many tables, limiting to first 100")
                    break

            return "\n\n".join(text_content)

        except ImportError:
            raise DocumentError(
                "DOCX processing library not available. Install python-docx."
            )
        except MemoryError:
            raise DocumentError("DOCX file too large to process - insufficient memory")
        except Exception as e:
            raise DocumentError(f"DOCX extraction failed: {e}")

    async def _extract_text(self, file_path: str) -> str:
        """Extract text from plain text files (TXT, MD, etc) with memory optimization."""
        try:
            # Check file size before loading
            file_size = os.path.getsize(file_path)
            if file_size > 100 * 1024 * 1024:  # 100MB limit for text files
                raise DocumentError("Text file too large (>100MB)")

            # Try multiple encodings
            encodings = ["utf-8", "utf-8-sig", "latin-1", "cp1252"]

            for encoding in encodings:
                try:
                    # For large files, read in chunks to avoid memory issues
                    if file_size > 10 * 1024 * 1024:  # 10MB
                        content_chunks = []
                        with open(file_path, "r", encoding=encoding) as file:
                            while True:
                                chunk = file.read(self.chunk_size)
                                if not chunk:
                                    break
                                content_chunks.append(chunk)
                                # Yield control periodically
                                if len(content_chunks) % 100 == 0:
                                    await asyncio.sleep(0)
                                    self._check_memory_usage()
                        return "".join(content_chunks)
                    else:
                        # Small files can be read normally
                        with open(file_path, "r", encoding=encoding) as file:
                            return file.read()
                except UnicodeDecodeError:
                    continue

            # If all encodings fail, try binary mode and decode with errors='ignore'
            with open(file_path, "rb") as file:
                raw_content = file.read()
                return raw_content.decode("utf-8", errors="ignore")

        except MemoryError:
            raise DocumentError("Text file too large to process - insufficient memory")
        except Exception as e:
            raise DocumentError(f"Text file extraction failed: {e}")

    async def _extract_text_streaming(self, file_path: str) -> AsyncIterator[str]:
        """Stream text content from plain text files."""
        try:
            encodings = ["utf-8", "utf-8-sig", "latin-1", "cp1252"]

            for encoding in encodings:
                try:
                    with open(file_path, "r", encoding=encoding) as file:
                        while True:
                            chunk = file.read(self.chunk_size)
                            if not chunk:
                                break
                            yield chunk
                            await asyncio.sleep(0)  # Allow other tasks
                    return  # Successfully processed
                except UnicodeDecodeError:
                    continue

            # Fallback to binary mode with error handling
            with open(file_path, "rb") as file:
                while True:
                    chunk = file.read(self.chunk_size)
                    if not chunk:
                        break
                    yield chunk.decode("utf-8", errors="ignore")
                    await asyncio.sleep(0)

        except Exception as e:
            raise DocumentError(f"Text streaming failed: {e}")

    async def _extract_rtf(self, file_path: str) -> str:
        """Extract text from RTF file."""
        try:
            from striprtf.striprtf import rtf_to_text

            with open(file_path, "r", encoding="utf-8") as file:
                rtf_content = file.read()

            return rtf_to_text(rtf_content)

        except ImportError:
            raise DocumentError(
                "RTF processing library not available. Install striprtf."
            )
        except Exception as e:
            raise DocumentError(f"RTF extraction failed: {e}")

    def get_file_info(self, file_path: str) -> Dict[str, Any]:
        """
        Get file information and metadata.

        Args:
            file_path: Path to the file

        Returns:
            dict: File information
        """
        try:
            file_path_obj = Path(file_path)

            # Get basic file info
            stat = file_path_obj.stat()
            mime_type, _ = mimetypes.guess_type(str(file_path_obj))

            return {
                "filename": file_path_obj.name,
                "size": stat.st_size,
                "extension": file_path_obj.suffix.lower().lstrip("."),
                "mime_type": mime_type,
                "created": stat.st_ctime,
                "modified": stat.st_mtime,
                "is_supported": file_path_obj.suffix.lower().lstrip(".")
                in self.supported_types,
            }

        except Exception as e:
            logger.error(f"Failed to get file info for {file_path}: {e}")
            return {"filename": Path(file_path).name, "error": str(e)}

    def validate_file(self, file_path: str, max_size: int = 10485760) -> bool:
        """
        Validate file for processing.

        Args:
            file_path: Path to the file
            max_size: Maximum file size in bytes

        Returns:
            bool: True if file is valid

        Raises:
            ValidationError: If file is invalid
        """
        file_path_obj = Path(file_path)

        if not file_path_obj.exists():
            raise ValidationError("File does not exist")

        if not file_path_obj.is_file():
            raise ValidationError("Path is not a file")

        # Check file size
        if file_path_obj.stat().st_size > max_size:
            raise ValidationError(
                f"File size exceeds maximum allowed ({max_size} bytes)"
            )

        # Check file type
        file_type = file_path_obj.suffix.lower().lstrip(".")
        if file_type not in self.supported_types:
            raise ValidationError(f"File type '{file_type}' is not supported")

        return True
